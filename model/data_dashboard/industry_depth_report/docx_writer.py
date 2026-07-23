# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLACK = RGBColor(0, 0, 0)


def ymd(date_value: str | None) -> str:
    if not date_value:
        return ""
    s = str(date_value).replace("-", "")
    return s[:8] if len(s) >= 8 else s


def pretty_date(s: str | None) -> str:
    t = ymd(s)
    if len(t) == 8:
        return f"{t[:4]}-{t[4:6]}-{t[6:8]}"
    return str(s or "")


def safe_float(v: Any, default: float | None = None) -> float | None:
    try:
        if v is None or v == "":
            return default
        x = float(v)
        if x != x:
            return default
        return x
    except Exception:
        return default


def fmt_num(v: Any, digits: int = 2) -> str:
    x = safe_float(v)
    return "-" if x is None else f"{x:.{digits}f}"


def fmt_pct_ratio(v: Any, digits: int = 1) -> str:
    x = safe_float(v)
    return "-" if x is None else f"{x * 100:.{digits}f}%"


def fmt_pct_value(v: Any, digits: int = 1) -> str:
    x = safe_float(v)
    return "-" if x is None else f"{x:.{digits}f}%"


def clean_text(text: Any) -> str:
    value = str(text or "")
    repl = {
        "数据来源": "",
        "数据源": "",
        "图表结论：": "",
        "图表结论": "",
        "本地数据库": "",
        "质量门": "",
        "API": "",
        "AI": "",
        "提示词": "",
        "行业定义": "",
        "若": "当",
        "如果": "当",
        "假设": "情形",
        "可能": "",
        "或许": "",
        "不确定": "",
        "不是而是": "",
        "、": "，",
        "“": "",
        "”": "",
        "\"": "",
        "?": "",
        "？": "。",
    }
    for old, new in repl.items():
        value = value.replace(old, new)
    return value.strip()


class WordWriter:
    def set_run_font(self, run, size: float = 10.5, bold: bool = False) -> None:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = BLACK

    def set_para(self, paragraph, first_line: bool = True, align=None) -> None:
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        if first_line:
            fmt.first_line_indent = Cm(0.74)
        if align is not None:
            paragraph.alignment = align

    def add_p(self, doc: Document, text: Any = "", size: float = 10.5, bold: bool = False, first_line: bool = True, align=None):
        value = clean_text(text)
        if not value:
            return None
        for line in value.splitlines():
            if not line.strip():
                continue
            p = doc.add_paragraph()
            self.set_para(p, first_line=first_line, align=align)
            r = p.add_run(line.strip())
            self.set_run_font(r, size=size, bold=bold)
        return None

    def add_h(self, doc: Document, text: str, level: int = 1):
        p = doc.add_paragraph()
        self.set_para(p, first_line=False)
        r = p.add_run(clean_text(text))
        self.set_run_font(r, size=15 if level == 1 else 12, bold=True)

    def set_cell(self, cell, text: Any, bold: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        self.set_para(p, first_line=False)
        r = p.add_run(clean_text(text))
        self.set_run_font(r, size=9.5, bold=bold)

    def add_table(self, doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
        if not rows:
            return
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            self.set_cell(table.rows[0].cells[i], h, True)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                self.set_cell(cells[i], value)
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_mar = OxmlElement("w:tcMar")
                for m in ["top", "left", "bottom", "right"]:
                    node = OxmlElement(f"w:{m}")
                    node.set(qn("w:w"), "70")
                    node.set(qn("w:type"), "dxa")
                    tc_mar.append(node)
                tc_pr.append(tc_mar)

    def add_fig(self, doc: Document, path: str | None, comment: Any = "") -> None:
        if path and Path(path).exists():
            p = doc.add_paragraph()
            self.set_para(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run().add_picture(path, width=Cm(15.8))
            self.add_p(doc, comment, size=10.5)

    def configure_doc(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
        for style_name in ["Normal", "Body Text"]:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
            style.font.size = Pt(10.5)
            style.font.color.rgb = BLACK
            style.paragraph_format.line_spacing = 1
            style.paragraph_format.space_after = Pt(0)

    def write(self, payload: dict[str, Any], out_dir: Path) -> Path:
        doc = Document()
        self.configure_doc(doc)
        industry = payload["industry"]
        figs = payload.get("figures", {})
        sections = payload.get("sections", {})
        m = payload.get("market", {})
        f = payload.get("financial", {})

        p = doc.add_paragraph()
        self.set_para(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(f"{industry}行业深度研究报告")
        self.set_run_font(r, size=18, bold=True)
        self.add_p(doc, f"报告日期：{pretty_date(payload.get('as_of'))}    回看区间：{pretty_date(payload.get('start_date'))} 至 {pretty_date(payload.get('as_of'))}", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)

        self.add_h(doc, "一 投资结论")
        self.add_p(doc, sections.get("investment_conclusion", ""))
        self.add_fig(doc, figs.get("thesis_map"), sections.get("thesis_map", ""))

        self.add_h(doc, "二 周期位置与市场定价")
        self.add_fig(doc, figs.get("relative_return"), sections.get("relative_return", ""))
        self.add_fig(doc, figs.get("drawdown"), sections.get("drawdown", ""))
        self.add_fig(doc, figs.get("valuation_band"), sections.get("valuation_band", ""))

        self.add_h(doc, "三 需求边际与外部环境")
        self.add_fig(doc, figs.get("demand_financial"), sections.get("demand_financial", ""))
        self.add_fig(doc, figs.get("macro_cycle"), sections.get("macro_cycle", ""))

        self.add_h(doc, "四 供给约束与资本周期")
        self.add_fig(doc, figs.get("supply_capital"), sections.get("supply_capital", ""))
        self.add_fig(doc, figs.get("dupont"), sections.get("dupont", ""))

        self.add_h(doc, "五 价格机制与利润弹性")
        self.add_fig(doc, figs.get("profit_pool"), sections.get("profit_pool", ""))
        self.add_fig(doc, figs.get("margin_roe"), sections.get("margin_roe", ""))

        self.add_h(doc, "六 竞争格局与财务分化")
        self.add_fig(doc, figs.get("financial_dispersion"), sections.get("financial_dispersion", ""))
        self.add_table(doc, ["指标", "最新读数", "含义"], [
            ["收入同比中位数", fmt_pct_value(f.get("median_tr_yoy")), "需求进入报表的力度"],
            ["净利润同比中位数", fmt_pct_value(f.get("median_netprofit_yoy")), "价格成本后的兑现"],
            ["毛利率中位数", fmt_pct_value(f.get("median_gross_margin")), "价格和成本关系"],
            ["净利率中位数", fmt_pct_value(f.get("median_net_margin")), "费用和经营杠杆"],
            ["ROE中位数", fmt_pct_value(f.get("median_roe")), "资本回报"],
            ["PE分位", fmt_pct_ratio(m.get("median_pe_pctile")), "估值赔率"],
            ["二十日资金净流入", f"{fmt_num(m.get('net_mf_20d_yi'), 1)}亿元", "交易边际"],
        ])

        self.add_h(doc, "七 资金边际与配置结论")
        self.add_fig(doc, figs.get("moneyflow_crowding"), sections.get("moneyflow_crowding", ""))
        self.add_p(doc, sections.get("final_view", ""))

        self.add_h(doc, "八 跟踪指标与反证")
        self.add_fig(doc, figs.get("tracking_table"), sections.get("tracking_table", ""))

        docx_path = out_dir / f"{industry}行业深度研究报告_{payload.get('as_of')}.docx"
        doc.save(docx_path)
        return docx_path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--payload", required=True)
    parser.add_argument("--out-dir", required=True)
    args = parser.parse_args()
    payload = json.loads(Path(args.payload).read_text(encoding="utf-8"))
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = WordWriter().write(payload, out_dir)
    print(json.dumps({"docx": str(path)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
